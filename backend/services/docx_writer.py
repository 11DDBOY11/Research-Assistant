"""
docx_writer.py — Convert Markdown paper + bibliography to DOCX.

Generates:
1. Paper sections (Introduction, Related Work, Methodology, Results, Conclusion)
2. References section — auto-generated from paper_metadata, sorted and
   deduplicated by paper_id actually cited in generated_sentences.
3. Transparency Report appendix inline in the same document.
"""
from __future__ import annotations

import io
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_docx(
    sections: dict[str, list[str]],    # {section_name: [sentence, ...]}
    bibliography: list[dict],           # [{title, authors, year, paper_id}, ...]
    transparency: dict,
    output_path: str | Path,
) -> None:
    """
    Write the full paper + transparency report to a DOCX file.

    sections: ordered dict of section_name → list of sentences
    bibliography: papers actually cited, sorted by first author surname then year
    transparency: the transparency report dict from Stage 7
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_heading("Generated Research Paper", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # ── Paper sections ────────────────────────────────────────────────────────
    SECTION_DISPLAY = {
        "introduction": "1. Introduction",
        "related_work": "2. Related Work",
        "methodology": "3. Methodology",
        "results": "4. Results and Discussion",
        "conclusion": "5. Conclusion",
    }

    for section_key, display_name in SECTION_DISPLAY.items():
        sentences = sections.get(section_key, [])
        if not sentences:
            continue
        doc.add_heading(display_name, level=1)
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        for sentence in sentences:
            run = para.add_run(sentence + " ")
            run.font.size = Pt(11)

    # ── References ────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("References", level=1)

    sorted_bib = sorted(
        bibliography,
        key=lambda p: (
            (p.get("authors") or [""])[0].split()[-1].lower() if p.get("authors") else "",
            p.get("year") or 9999,
        ),
    )

    for i, paper in enumerate(sorted_bib, 1):
        authors = paper.get("authors") or []
        author_str = ", ".join(authors) if authors else "Unknown Author"
        year = paper.get("year") or "n.d."
        title = paper.get("title") or "Untitled"

        ref_para = doc.add_paragraph(style="List Number")
        ref_para.clear()
        run = ref_para.add_run(f"{author_str} ({year}). ")
        run.font.size = Pt(10)
        run.bold = True
        run2 = ref_para.add_run(f"{title}.")
        run2.font.size = Pt(10)
        run2.italic = True

    # ── Transparency Report appendix ──────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Transparency Report", level=1)
    _write_transparency_section(doc, transparency)

    doc.save(str(output_path))


def _write_transparency_section(doc: Document, report: dict) -> None:
    stats = [
        ("Papers analyzed", report.get("papers_analyzed", 0)),
        ("Claims extracted", report.get("claims_extracted", 0)),
        ("Sentences generated", report.get("sentences_generated", 0)),
        ("Sentences verified ✓", report.get("sentences_verified", 0)),
        ("Sentences rejected ✗", report.get("sentences_rejected", 0)),
        ("Open clarification requests", report.get("open_clarifications", 0)),
    ]

    doc.add_heading("Summary Statistics", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Count"
    for label, count in stats:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(count)

    # LLM usage / cost
    usage = report.get("llm_usage", {})
    if usage:
        doc.add_heading("API Usage & Estimated Cost", level=2)
        doc.add_paragraph(
            f"Total prompt tokens: {usage.get('total_prompt_tokens', 0):,}\n"
            f"Total completion tokens: {usage.get('total_completion_tokens', 0):,}\n"
            f"Estimated total cost: ${usage.get('total_cost_usd', 0.0):.4f} USD"
        )

    # Rejected sentences log
    rejected = report.get("rejected_details", [])
    if rejected:
        doc.add_heading("Rejected Sentences Log", level=2)
        for item in rejected:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item.get("text", "")[:200])
            run.font.size = Pt(9)
            reason_run = p.add_run(f"\n  Reason: {item.get('rejection_reason', 'unknown')}")
            reason_run.font.size = Pt(9)
            reason_run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)


def paper_to_markdown(
    sections: dict[str, list[str]],
    bibliography: list[dict],
) -> str:
    """Generate a plain Markdown version of the paper (for preview in UI)."""
    lines = ["# Generated Research Paper\n"]

    SECTION_DISPLAY = {
        "introduction": "## 1. Introduction",
        "related_work": "## 2. Related Work",
        "methodology": "## 3. Methodology",
        "results": "## 4. Results and Discussion",
        "conclusion": "## 5. Conclusion",
    }

    for key, heading in SECTION_DISPLAY.items():
        sentences = sections.get(key, [])
        if not sentences:
            continue
        lines.append(f"\n{heading}\n")
        lines.append(" ".join(sentences))
        lines.append("")

    lines.append("\n## References\n")
    sorted_bib = sorted(
        bibliography,
        key=lambda p: (
            (p.get("authors") or [""])[0].split()[-1].lower() if p.get("authors") else "",
            p.get("year") or 9999,
        ),
    )
    for i, paper in enumerate(sorted_bib, 1):
        authors = ", ".join(paper.get("authors") or ["Unknown"])
        year = paper.get("year") or "n.d."
        title = paper.get("title") or "Untitled"
        lines.append(f"{i}. **{authors}** ({year}). *{title}*.")

    return "\n".join(lines)
